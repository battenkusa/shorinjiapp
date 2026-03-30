import streamlit as st
from pathlib import Path
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage
import os
import json
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
import pdfplumber
from docx import Document
import time

st.set_page_config(page_title="葛飾区少林寺拳法連盟AIチーム", layout="wide")
st.title("🤝 葛飾区少林寺拳法連盟AIチーム")

llm = ChatOllama(model="llama3.2")

# 提出資料監視機能のヘルパー関数
def get_last_scan_time():
    """最後のスキャン時刻を取得"""
    scan_file = Path("output/last_scan.json")
    if scan_file.exists():
        try:
            with open(scan_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return datetime.fromisoformat(data.get('last_scan', '2000-01-01T00:00:00'))
        except:
            pass
    return datetime(2000, 1, 1)

def save_last_scan_time():
    """現在時刻をスキャン時刻として保存"""
    Path("output").mkdir(exist_ok=True)
    scan_file = Path("output/last_scan.json")
    with open(scan_file, 'w', encoding='utf-8') as f:
        json.dump({"last_scan": datetime.now().isoformat()}, f, ensure_ascii=False)

def read_document(file_path):
    """ドキュメントファイルを読み込む"""
    file_path = Path(file_path)
    if not file_path.exists():
        return ""
    
    try:
        if file_path.suffix.lower() == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif file_path.suffix.lower() == '.docx':
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_path.suffix.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        st.error(f"ファイル読み込みエラー: {file_path.name} - {str(e)}")
    
    return ""

def analyze_document(content, filename):
    """ドキュメント内容をOllamaで分析"""
    prompt = f"""以下の資料を分析して、次の項目を抽出してください：

1. 概要（どのような資料か）
2. 実施内容（何をする必要があるか）
3. 期限（提出期限や実施期限があれば具体的な日付）

資料ファイル名: {filename}
資料内容:
{content[:2000]}

分析結果を以下の形式で出力してください：
【概要】
（概要を記載）

【実施内容】
（実施内容を記載）

【期限】
（期限を記載、見つからない場合は「不明」）"""

    try:
        res = llm.invoke([HumanMessage(content=prompt)])
        return res.content
    except Exception as e:
        return f"分析エラー: {str(e)}"

def send_notification_email(subject, body):
    """Gmail経由でメール送信"""
    gmail_address = os.getenv('GMAIL_ADDRESS')
    gmail_password = os.getenv('GMAIL_APP_PASSWORD')
    
    if not gmail_address or not gmail_password:
        st.error("メール設定が不完全です。環境変数 GMAIL_ADDRESS と GMAIL_APP_PASSWORD を設定してください。")
        return False
    
    try:
        msg = MIMEMultipart()
        msg['From'] = gmail_address
        msg['To'] = "battenkusayokayoka@gmail.com"
        msg['Subject'] = subject
        
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(gmail_address, gmail_password)
        server.send_message(msg)
        server.quit()
        
        return True
    except Exception as e:
        st.error(f"メール送信エラー: {str(e)}")
        return False

def scan_submission_folder():
    """提出資料フォルダをスキャンして新しいファイルを検出"""
    folder_path = Path(r"C:\Users\batte\OneDrive\少林寺\葛飾区連盟\AI作業フォルダ\提出関連資料")
    
    # Windowsパスが存在しない場合は、テスト用のローカルパスを使用
    if not folder_path.exists():
        folder_path = Path("input")
        folder_path.mkdir(exist_ok=True)
    
    last_scan = get_last_scan_time()
    new_files = []
    
    if folder_path.exists():
        for file_path in folder_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in ['.pdf', '.docx', '.txt']:
                # ファイルの更新時刻をチェック
                file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                if file_time > last_scan:
                    new_files.append(file_path)
    
    return new_files, folder_path

# Streamlit起動時の自動スキャン
if 'startup_scan_done' not in st.session_state:
    with st.spinner("起動時スキャンを実行中..."):
        try:
            new_files, folder_path = scan_submission_folder()
            if new_files:
                st.sidebar.success(f"新しい資料 {len(new_files)} 件を検出しました")
                for file_path in new_files:
                    content = read_document(file_path)
                    if content:
                        analysis = analyze_document(content, file_path.name)
                        
                        # メール本文を生成
                        email_body = f"""新しい提出資料が検出されました。

ファイル名: {file_path.name}
検出時刻: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
ファイルパス: {file_path}

【AI分析結果】
{analysis}

【依頼事項】
上記資料の確認と必要な対応をお願いいたします。
期限がある場合は、期限の10日前までに内容を整理して対応準備を進めてください。

このメールは葛飾区少林寺拳法連盟AIシステムから自動送信されました。
"""
                        
                        # メール送信
                        subject = f"【提出資料監視】新資料検出: {file_path.name}"
                        send_notification_email(subject, email_body)
                
                save_last_scan_time()
            else:
                st.sidebar.info("新しい資料は見つかりませんでした")
        except Exception as e:
            st.sidebar.error(f"スキャンエラー: {str(e)}")
    
    st.session_state.startup_scan_done = True

menu = st.sidebar.radio("機能メニュー", [
    "📋 議事録作成",
    "✉️ 連絡文作成",
    "📢 お知らせ作成",
    "📅 活動計画作成",
    "📄 大会資料作成",
    "📬 提出資料監視",
])

# ── 議事録作成 ──────────────────────────
if menu == "📋 議事録作成":
    st.subheader("議事録作成")
    col1, col2 = st.columns(2)
    title   = col1.text_input("会議名", placeholder="例：3月定例ミーティング")
    date    = col2.text_input("日時",   placeholder="例：2025年3月28日 19:00〜")
    members = st.text_input("参加者",   placeholder="例：田中・鈴木・佐藤（計3名）")
    memo    = st.text_area("会議メモ", height=150,
                placeholder="・文化祭の出し物をカフェに決定\n・担当は田中・鈴木・佐藤\n・次回は4月5日集合")

    if st.button("議事録を生成する"):
        with st.spinner("AIが生成中..."):
            prompt = f"""以下の会議メモから、見やすい議事録を日本語で作成してください。
会議名：{title} / 日時：{date} / 参加者：{members}
メモ：{memo}"""
            res = llm.invoke([HumanMessage(content=prompt)])
        st.success("完成しました！")
        st.markdown(res.content)
        st.download_button("テキストで保存", res.content, "議事録.txt")

# ── 連絡文作成 ──────────────────────────
elif menu == "✉️ 連絡文作成":
    st.subheader("連絡文作成")
    col1, col2 = st.columns(2)
    to      = col1.text_input("宛先",   placeholder="例：部員の皆さん")
    sender  = col2.text_input("差出人", placeholder="例：部長 田中")
    tone    = st.selectbox("トーン", ["丁寧・敬語", "やや砕けた", "フレンドリー"])
    content = st.text_area("伝えたい内容", height=120,
                placeholder="例：次の練習は4月5日に変更。場所は第2体育館。")

    if st.button("連絡文を生成する"):
        with st.spinner("AIが生成中..."):
            prompt = f"""以下の情報をもとに、{tone}なトーンでメール文を日本語で作成してください。
宛先：{to} / 差出人：{sender}
内容：{content}"""
            res = llm.invoke([HumanMessage(content=prompt)])
        st.success("完成しました！")
        st.markdown(res.content)
        st.download_button("テキストで保存", res.content, "連絡文.txt")

# ── お知らせ作成 ────────────────────────
elif menu == "📢 お知らせ作成":
    st.subheader("お知らせ作成")
    usage   = st.radio("用途", ["掲示板・印刷", "LINE / SNS", "チラシ"], horizontal=True)
    content = st.text_area("お知らせ内容", height=120,
                placeholder="例：5月の大会でカフェを出します。メンバー募集中です。")

    if st.button("お知らせを生成する"):
        with st.spinner("AIが生成中..."):
            prompt = f"""{usage}向けのお知らせ文を日本語で作成してください。
内容：{content}"""
            res = llm.invoke([HumanMessage(content=prompt)])
        st.success("完成しました！")
        st.markdown(res.content)
        st.download_button("テキストで保存", res.content, "お知らせ.txt")

# ── 活動計画作成 ────────────────────────
elif menu == "📅 活動計画作成":
    st.subheader("活動計画作成")
    goal    = st.text_input("目標・イベント名", placeholder="例：春季区民体育大会")
    col1, col2 = st.columns(2)
    start   = col1.text_input("開始日", placeholder="例：4月1日")
    end     = col2.text_input("締切日", placeholder="例：5月31日")
    tasks   = st.text_area("やること（わかっているもの）",
                placeholder="例：参加者募集、会場準備、役割分担、リハーサル")

    if st.button("計画を生成する"):
        with st.spinner("AIが生成中..."):
            prompt = f"""以下の情報をもとに、週単位の活動スケジュールを日本語で作成してください。
目標：{goal} / 期間：{start}〜{end}
やること：{tasks}"""
            res = llm.invoke([HumanMessage(content=prompt)])
        st.success("完成しました！")
        st.markdown(res.content)
        st.download_button("テキストで保存", res.content, "活動計画.txt")

# ── 大会資料作成 ────────────────────────
elif menu == "📄 大会資料作成":
    st.subheader("大会資料作成（PDF → Word）")
    st.caption("昨年のPDFを読み込んで、今年版のWordファイルを自動生成します")

    pdf_path = st.text_input(
        "元になるPDFのパス",
        value=r"C:\Users\batte\OneDrive\少林寺\葛飾区連盟\2025年度\春季区大会\2025年度春季区民体育大会案rev２.pdf"
    )

    col1, col2, col3 = st.columns(3)
    year    = col1.text_input("今年度", value="2026")
    date    = col2.text_input("実施日", value="5/31")
    members = col3.text_input("人数",   value="80")

    out_path = st.text_input(
        "保存先（Wordファイル）",
        value=r"C:\Users\batte\myagent\output\2026年度春季区大会資料.docx"
    )

    if st.button("Word資料を生成する"):
        import pdfplumber
        from docx import Document as DocxDocument

        # ① PDFを読み込む
        with st.spinner("PDFを読み込み中..."):
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    pdf_text = "\n".join(
                        page.extract_text() for page in pdf.pages
                        if page.extract_text()
                    )
                st.info(f"PDF読み込み完了（{len(pdf_text)}文字）")
            except Exception as e:
                st.error(f"PDF読み込みエラー: {e}")
                st.stop()

        # ② AIに内容を更新させる
        with st.spinner("AIが内容を更新中..."):
            prompt = f"""以下は2025年度の大会資料のテキストです。
これをもとに、下記の変更点を反映した{year}年度版の資料文章を作成してください。

【変更点】
- 年度：2025年度 → {year}年度
- 実施日：{date}
- 参加人数：{members}人
- その他の日付・年度の記載も{year}年度に合わせて更新

【元の資料】
{pdf_text[:3000]}

【出力形式】
タイトル、目的、実施日、会場、参加者、スケジュール、注意事項などの項目を
整理した資料文章を日本語で出力してください。"""

            res = llm.invoke([HumanMessage(content=prompt)])
            updated_text = res.content

        # ③ Wordファイルに書き出す
        with st.spinner("Wordファイルを作成中..."):
            doc = DocxDocument()
            doc.add_heading(f"{year}年度 春季区民体育大会資料", level=1)

            for line in updated_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("【") or line.startswith("■") or line.endswith("】"):
                    doc.add_heading(line, level=2)
                else:
                    doc.add_paragraph(line)

            Path(out_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(out_path)

        st.success("完成しました！")
        st.code(out_path)
        st.subheader("生成された内容（プレビュー）")
        st.markdown(updated_text)

# ── 提出資料監視 ────────────────────────
elif menu == "📬 提出資料監視":
    st.subheader("提出資料監視・メール送信")
    st.caption("OneDriveの提出関連資料フォルダを監視し、新しいファイルが追加されたら自動で分析・メール送信します")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.info("📁 監視対象フォルダ")
        folder_path = Path(r"C:\Users\batte\OneDrive\少林寺\葛飾区連盟\AI作業フォルダ\提出関連資料")
        if not folder_path.exists():
            folder_path = Path("input")
        st.code(str(folder_path))
        
        if folder_path.exists():
            file_count = len([f for f in folder_path.rglob('*') if f.is_file() and f.suffix.lower() in ['.pdf', '.docx', '.txt']])
            st.metric("監視対象ファイル数", f"{file_count}件")
        else:
            st.error("監視フォルダが見つかりません")
    
    with col2:
        st.info("📧 メール送信設定")
        gmail_address = os.getenv('GMAIL_ADDRESS')
        gmail_password = os.getenv('GMAIL_APP_PASSWORD')
        
        if gmail_address and gmail_password:
            st.success("✅ メール設定完了")
            st.text(f"送信元: {gmail_address}")
            st.text("送信先: battenkusayokayoka@gmail.com")
        else:
            st.error("❌ メール設定が不完全")
            st.text("環境変数を設定してください:")
            st.code("GMAIL_ADDRESS\nGMAIL_APP_PASSWORD")
    
    st.divider()
    
    col1, col2 = st.columns(2)
    
    # 最後のスキャン時刻表示
    with col1:
        last_scan = get_last_scan_time()
        st.info("📅 最後のスキャン")
        st.text(last_scan.strftime('%Y年%m月%d日 %H:%M:%S'))
    
    # 手動スキャンボタン
    with col2:
        if st.button("🔍 手動スキャン実行", type="primary"):
            with st.spinner("スキャンを実行中..."):
                try:
                    new_files, scan_folder = scan_submission_folder()
                    
                    if new_files:
                        st.success(f"新しい資料 {len(new_files)} 件を検出しました")
                        
                        for i, file_path in enumerate(new_files):
                            st.subheader(f"📄 {file_path.name}")
                            
                            # ファイル情報
                            file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                            st.text(f"更新日時: {file_time.strftime('%Y年%m月%d日 %H:%M:%S')}")
                            
                            # ファイル内容読み込み
                            with st.spinner(f"ファイル {i+1}/{len(new_files)} を分析中..."):
                                content = read_document(file_path)
                                
                                if content:
                                    # AI分析
                                    analysis = analyze_document(content, file_path.name)
                                    st.markdown("**AI分析結果:**")
                                    st.markdown(analysis)
                                    
                                    # メール本文生成
                                    email_body = f"""新しい提出資料が検出されました。

ファイル名: {file_path.name}
検出時刻: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
ファイルパス: {file_path}

【AI分析結果】
{analysis}

【依頼事項】
上記資料の確認と必要な対応をお願いいたします。
期限がある場合は、期限の10日前までに内容を整理して対応準備を進めてください。

このメールは葛飾区少林寺拳法連盟AIシステムから自動送信されました。
"""
                                    
                                    # メール送信
                                    subject = f"【提出資料監視】新資料検出: {file_path.name}"
                                    if send_notification_email(subject, email_body):
                                        st.success(f"✅ {file_path.name} の通知メールを送信しました")
                                    else:
                                        st.error(f"❌ {file_path.name} のメール送信に失敗しました")
                                else:
                                    st.warning(f"⚠️ {file_path.name} の内容を読み込めませんでした")
                        
                        save_last_scan_time()
                        
                    else:
                        st.info("新しい資料は見つかりませんでした")
                        
                except Exception as e:
                    st.error(f"スキャンエラー: {str(e)}")
    
    st.divider()
    
    # 設定情報
    with st.expander("⚙️ 設定情報"):
        st.markdown("""
        **監視対象ファイル形式:**
        - PDF (.pdf)
        - Microsoft Word (.docx)  
        - テキストファイル (.txt)
        
        **自動実行:**
        - Streamlit起動時に自動スキャン実行
        - 前回スキャン時刻以降に更新されたファイルを検出
        
        **AI分析項目:**
        - 概要（どのような資料か）
        - 実施内容（何をする必要があるか）
        - 期限（提出期限や実施期限）
        
        **メール送信:**
        - Gmail SMTP経由で送信
        - 送信先: battenkusayokayoka@gmail.com
        - 件名: 【提出資料監視】新資料検出: ファイル名
        """)